import os
import re
import logging
import time
import signal
from typing import List, Dict, Any, Optional, Tuple, Union
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, TimeoutError

# Import LangChain document types and text splitters
try:
    from langchain_core.documents import Document as LangchainDocument
    from langchain.text_splitter import RecursiveCharacterTextSplitter
except ImportError:
    try:
        from langchain.schema import Document as LangchainDocument
        from langchain.text_splitter import RecursiveCharacterTextSplitter
    except ImportError:
        raise ImportError("langchain-core or langchain is required for document processing")

# Import third-party libraries with error handling
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None
    
try:
    import pdfplumber
except ImportError:
    pdfplumber = None
    
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None
    
try:
    import pandas as pd
except ImportError:
    pd = None

# Configure logging
logger = logging.getLogger(__name__)


class DocumentLoader:
    """
    A document loader that handles various file formats with enhanced table extraction.
    """
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._load_pdf,
            '.docx': self._load_docx,
            '.txt': self._load_text,
            '.md': self._load_text,
            '.csv': self._load_text,
            '.xlsx': self._load_excel_with_tables,
            '.xls': self._load_excel_with_tables,
        }
    
    def load_document(self, file_path: str) -> List[LangchainDocument]:
        """
        Load a document from the given file path.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of LangchainDocument objects
        """
        file_path = str(file_path)
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext in self.supported_formats:
            return self.supported_formats[file_ext](file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _fallback_pdf_extraction(self, file_path: str) -> str:
        """
        Fallback method to extract text from PDF when primary methods fail.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text as a string
        """
        logger.warning(f"Attempting fallback PDF extraction for: {file_path}")
        
        # Try PyMuPDF as a fallback
        if fitz:
            try:
                text = []
                with fitz.open(file_path) as doc:
                    for page in doc:
                        text.append(page.get_text())
                return "\n\n".join(text)
            except Exception as e:
                logger.warning(f"Fallback PyMuPDF extraction failed: {str(e)}")
        
        # As a last resort, try to use the system's pdftotext if available
        try:
            import subprocess
            result = subprocess.run(
                ['pdftotext', '-layout', file_path, '-'],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout
        except (FileNotFoundError, subprocess.SubprocessError, Exception) as e:
            logger.warning(f"pdftotext extraction failed: {str(e)}")
        
        return ""
    
    def _process_extracted_text(self, text: str, file_path: str) -> List[LangchainDocument]:
        """
        Process extracted text into LangchainDocument objects.
        
        Args:
            text: Extracted text content
            file_path: Source file path for metadata
            
        Returns:
            List of LangchainDocument objects
        """
        try:
            # Clean the extracted text
            cleaned_text = clean_text(text)
            
            # Create a metadata dictionary
            metadata = {
                'source': str(file_path),
                'file_name': os.path.basename(file_path),
                'file_type': 'pdf',
                'extraction_method': 'pdfplumber with fallbacks'
            }
            
            # Create a single document with the extracted text
            return [LangchainDocument(
                page_content=cleaned_text,
                metadata=metadata
            )]
            
        except Exception as e:
            logger.error(f"Error processing extracted text: {str(e)}")
            raise
    
    def _load_pdf(self, file_path: str) -> List[LangchainDocument]:
        """
        Load a PDF file with enhanced table extraction.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            List of LangchainDocument objects
            
        Raises:
            ImportError: If required dependencies are missing
            ValueError: If no text could be extracted from the PDF
        """
        if not pdfplumber:
            raise ImportError("pdfplumber is required for PDF processing")
        
        try:
            logger.info(f"Extracting text from PDF: {file_path}")
            text = extract_text_from_pdf(file_path, max_pages=5)  # Limit to 5 pages by default
            
            if not text or not text.strip():
                logger.warning("No text could be extracted from the PDF")
                # Try fallback extraction method
                text = self._fallback_pdf_extraction(file_path)
                
            if not text or not text.strip():
                raise ValueError("No text could be extracted from the PDF after fallback methods")
                
            return self._process_extracted_text(text, file_path)
            
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {str(e)}")
            raise
    
    def _load_docx(self, file_path: str) -> List[LangchainDocument]:
        """Load a DOCX file."""
        if not DocxDocument:
            raise ImportError("python-docx is required for DOCX processing")
            
        try:
            text = extract_text_from_docx(file_path)
            return [LangchainDocument(
                page_content=text,
                metadata={"source": file_path, "file_type": "docx"}
            )]
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {str(e)}")
            raise
    
    def _load_text(self, file_path: str) -> List[LangchainDocument]:
        """Load a plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
                
            return [LangchainDocument(
                page_content=text,
                metadata={"source": file_path, "file_type": "text"}
            )]
        except Exception as e:
            logger.error(f"Error loading text file {file_path}: {str(e)}")
            raise
    
    def _load_excel_with_tables(self, file_path: str) -> List[LangchainDocument]:
        """Load an Excel file and convert each sheet to a markdown table."""
        if not pd:
            raise ImportError("pandas is required for Excel file processing")
            
        try:
            xl = pd.ExcelFile(file_path)
            docs = []
            
            for sheet_name in xl.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convert DataFrame to markdown
                markdown_table = df.to_markdown(index=False)
                
                # Create a document for each sheet
                docs.append(LangchainDocument(
                    page_content=f"# {sheet_name}\n\n{markdown_table}",
                    metadata={
                        "source": file_path,
                        "file_type": "excel",
                        "sheet_name": sheet_name
                    }
                ))
                
            return docs
                
        except Exception as e:
            logger.error(f"Error loading Excel file {file_path}: {str(e)}")
            raise

# Common insurance-related terms to help identify insurance documents
INSURANCE_KEYWORDS = {
    'insurance', 'policy', 'premium', 'coverage', 'benefit', 'claim', 'deductible',
    'sum insured', 'policyholder', 'insured', 'insurer', 'exclusion', 'endorsement',
    'rider', 'annuity', 'maturity', 'surrender', 'nominee', 'proposer', 'underwriting'
}

def is_insurance_document(text: str, threshold: int = 3) -> bool:
    """
    Check if the given text appears to be an insurance-related document.
    
    Args:
        text: The text content to analyze
        threshold: Minimum number of insurance-related keywords to match
        
    Returns:
        bool: True if the text appears to be insurance-related, False otherwise
    """
    if not text:
        return False
    
    # Convert text to lowercase for case-insensitive matching
    text_lower = text.lower()
    
    # Count how many insurance keywords appear in the text
    matches = sum(1 for keyword in INSURANCE_KEYWORDS if keyword in text_lower)
    
    return matches >= threshold

def extract_tables_with_camelot(pdf_path: str, page_num: int) -> str:
    """
    Extract tables from a PDF page using camelot.
    
    Args:
        pdf_path: Path to the PDF file
        page_num: Page number (1-based index)
        
    Returns:
        str: Extracted tables in markdown format
    """
    if not camelot:
        return ""
        
    try:
        # Try to extract tables using camelot
        tables = camelot.read_pdf(
            pdf_path, 
            pages=str(page_num),
            flavor='lattice',  # Try lattice first (for tables with lines)
            strip_text='\n',
            suppress_stdout=True
        )
        
        # If no tables found with lattice, try stream
        if not tables:
            tables = camelot.read_pdf(
                pdf_path,
                pages=str(page_num),
                flavor='stream',  # Try stream (for tables without lines)
                strip_text='\n',
                suppress_stdout=True
            )
        
        result = []
        for i, table in enumerate(tables):
            # Convert to markdown
            df = table.df
            if not df.empty:
                result.append(f"\n```table\nTable {i+1} (Confidence: {table.parsing_report['accuracy']:.1f}%)\n")
                result.append(df.to_markdown(index=False))
                result.append("\n```\n")
        
        return "".join(result)
    except Exception as e:
        logger.warning(f"Error extracting tables with camelot: {str(e)}")
        return ""

def extract_excluded_items(text: str) -> str:
    """
    Extract excluded items from policy text using pattern matching.
    
    This function looks for common patterns in insurance policy documents that
    indicate a list of excluded items, such as "List I - Items not covered".
    It handles both numbered and bulleted lists and formats them in markdown.
    
    Args:
        text: The text content to search for excluded items
        
    Returns:
        str: Formatted excluded items in markdown, or empty string if none found
    """
    # Common patterns that might indicate the start of an exclusion list
    exclusion_patterns = [
        # Matches "List I - Items not covered" with variations
        r'(?i)(List\s*I[\s\-]*(?:Items?\s*(?:not\s*covered|excluded)|[^\n]*))[\s\S]*?(?=\n\s*List\s*[I|V]|\n\s*\d|\n\s*[A-Z][^\n]*\n\s*[A-Z]|$)',
        # Matches "Exclusions:" followed by bullet points or numbered lists
        r'(?i)(Exclusions?\s*:[\s\S]*?)(?=\n\s*[A-Z][^\n]*\n\s*[A-Z]|\n\s*\d|$)',
        # Matches "Items not covered:" with various list formats
        r'(?i)(Items?\s*not\s*covered\s*:[\s\S]*?)(?=\n\s*[A-Z][^\n]*\n\s*[A-Z]|\n\s*\d|$)',
        # Matches "The following are excluded:" patterns
        r'(?i)(The\s+following\s+(?:items\s+)?(?:are|is)\s+excluded\s*:[\s\S]*?)(?=\n\s*[A-Z][^\n]*\n\s*[A-Z]|\n\s*\d|$)',
        # Matches any text with common exclusion keywords followed by a list
        r'(?i)((?:Exclusions?|Items?\s+not\s+covered|Not\s+covered|Excluded\s+items?)[\s\S]*?\n(?:\s*[•\-*]\s*[^\n]*\n)+)',
        # Special case for HDFC policy format
        r'(?i)(List\s*I[\s\-]*Items?[^\n]*\n[\s\S]*?)(?=\n\s*List\s*II|\n\s*[A-Z][^\n]*\n\s*[A-Z]|$)'
    ]
    
    # Additional patterns for common exclusion list formats
    list_patterns = [
        # Numbered lists (1., 2., etc.)
        r'(?m)^\s*\d+[.)]\s+.*(?:\n\s+.*)*',
        # Bullet points (•, -, *)
        r'(?m)^\s*[•\-*]\s+.*(?:\n\s+.*)*',
        # Lettered lists (a., b., etc.)
        r'(?m)^\s*[a-z][.)]\s+.*(?:\n\s+.*)*',
        # Roman numerals (i., ii., etc.)
        r'(?m)^\s*[ivx]+[.)]\s+.*(?:\n\s+.*)*',
        # Lines starting with common exclusion keywords
        r'(?m)^\s*(?:Exclusion|Not\s+covered|Excluded|Not\s+eligible|Not\s+payable|Not\s+included|Not\s+applicable)[^\n]*',
        # Lines with specific exclusion terms
        r'(?i)(?:\b(?:exclusion|not\s+covered|excluded|not\s+eligible|not\s+payable|not\s+included|not\s+applicable)\b[^\n]*)',
        # Lines that look like exclusions (e.g., "- Baby food")
        r'(?m)^\s*([^\n]+:\s*\n(?:\s+[^\n]+\n?)+)'
    ]
    
    # Common exclusion keywords to look for
    exclusion_keywords = [
        'exclu', 'not covered', 'not pay', 'not eligible', 'not included',
        'not applicable', 'not covered', 'not covered', 'not covered',
        'baby food', 'baby', 'food', 'nutrition', 'diet', 'formula', 'powder'
    ]
    
    # Try to find the exclusion list using the patterns
    for pattern in exclusion_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE | re.DOTALL)
        for match in matches:
            # Extract the matched text and clean it up
            excluded_text = match.group(1) if len(match.groups()) > 0 else match.group(0)
            excluded_text = excluded_text.strip()
            
            # If we found a potential exclusion section, try to extract the list items
            if excluded_text:
                # Look for list items within the matched section
                list_items = []
                for list_pattern in list_patterns:
                    list_matches = re.findall(list_pattern, excluded_text, re.IGNORECASE | re.MULTILINE)
                    if list_matches:
                        # Flatten the list of matches (some patterns may return tuples)
                        for item in list_matches:
                            if isinstance(item, tuple):
                                # Take the first non-empty group
                                item = next((x for x in item if x), "")
                            if item and item.strip():
                                list_items.append(item.strip())
                
                # If we found list items, format them nicely
                if list_items:
                    # Clean up list items and filter by relevance
                    cleaned_items = []
                    for item in list_items:
                        # Remove any leading numbers or bullets
                        item = re.sub(r'^\s*[\d•\-*]+[.)]?\s*', '', item)
                        # Clean up any remaining whitespace
                        item = ' '.join(item.split())
                        # Only include items that are relevant or contain exclusion keywords
                        if item and (len(item) > 10 or any(keyword in item.lower() for keyword in exclusion_keywords)):
                            cleaned_items.append(f"- {item}")
                    
                    # If we have cleaned items, return them as a formatted markdown list
                    if cleaned_items:
                        # Get the section title if available
                        title_match = re.search(r'^[^:\n]+:', excluded_text, re.IGNORECASE)
                        title = title_match.group(0).strip() if title_match else "## Items Not Covered"
                        
                        # Add a note about the source of this information
                        source_note = "\n\n*This information was extracted from the policy document's exclusion section.*"
                        return f"{title}\n\n" + '\n'.join(cleaned_items) + source_note
                
                # If we didn't find list items but have a good match, return the raw text
                return f"## Items Not Covered\n\n{excluded_text}"
    
    # If no exclusions found with patterns, try to find any list that might contain exclusions
    for list_pattern in list_patterns:
        list_matches = re.findall(list_pattern, text, re.IGNORECASE | re.MULTILINE)
        if list_matches:
            # Filter for items that contain exclusion-related keywords
            exclusion_items = []
            for item in list_matches:
                if isinstance(item, tuple):
                    item = next((x for x in item if x), "")
                item = item.strip()
                if item and any(keyword in item.lower() for keyword in exclusion_keywords):
                    # Clean up the item
                    item = re.sub(r'^\s*[\d•\-*]+[.)]?\s*', '', item)
                    item = ' '.join(item.split())
                    if item and len(item) > 5:  # Lower the length threshold for potential exclusions
                        exclusion_items.append(f"- {item}")
            
            if exclusion_items:
                return "## Potential Exclusions Found\n\n" + '\n'.join(exclusion_items) + "\n\n*This information was automatically extracted and may be incomplete.*"
    
    # If we still haven't found anything, try to find any mention of exclusions
    exclusion_mentions = re.findall(r'(?i)(?:exclusion|not covered|excluded)[^\n]{10,200}', text)
    if exclusion_mentions:
        return "## Exclusion Information\n\n" + "\n- ".join(["Mentions of exclusions:"] + [m.strip() for m in exclusion_mentions if m.strip()])
    
    return ""

def clean_cell(cell):
    """Clean and format a single table cell or list item."""
    if cell is None:
        return ""
    # Convert to string, remove extra whitespace, and clean up
    text = str(cell).strip()
    # Replace multiple spaces with single space
    text = ' '.join(text.split())
    # Remove any remaining control characters except newlines and tabs
    text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t')
    return text

def format_table_as_markdown(rows):
    """Format a 2D array of cells as a markdown table."""
    if not rows or not any(rows):
        return ""
        
    # Clean all cells
    cleaned_rows = [[clean_cell(cell) for cell in row] for row in rows]
    
    # Remove empty rows
    cleaned_rows = [row for row in cleaned_rows if any(cell for cell in row)]
    
    if not cleaned_rows:
        return ""
        
    # Format as markdown table
    headers = cleaned_rows[0]
    separator = ['---'] * len(headers)
    
    # Build the markdown table
    table = [
        '| ' + ' | '.join(headers) + ' |',
        '|' + '|'.join(separator) + '|'
    ]
    
    for row in cleaned_rows[1:]:
        table.append('| ' + ' | '.join(row) + ' |')
    
    return '\n'.join(table) + '\n\n'

def extract_tables_with_pdfplumber(page) -> str:
    """
    Extract tables from a PDF page using pdfplumber with improved handling.
    
    This function uses multiple strategies to extract tables and lists from PDFs,
    including both bordered and borderless tables, and various list formats.
    
    Args:
        page: pdfplumber page object
        
    Returns:
        str: Extracted tables and lists in markdown format with improved formatting
    """
    if not pdfplumber:
        return ""
    
    # Initialize a set to track seen tables to avoid duplicates
    seen_tables = set()
    all_tables = []
    
    # Table extraction strategies - try multiple approaches
    strategies = [
        # Default strategy for regular tables
        {},
        
        # Strategy for borderless tables
        {
            'vertical_strategy': 'text',
            'horizontal_strategy': 'text',
            'keep_blank_chars': True,
            'edge_min_length': 3,
            'min_words_vertical': 1
        },
        
        # Strategy for tables with explicit borders
        {
            'vertical_strategy': 'lines',
            'horizontal_strategy': 'lines',
            'snap_tolerance': 3,
            'join_tolerance': 3,
            'edge_min_length': 3,
            'min_words_vertical': 1
        }
    ]
    
    # First, try to extract structured tables
    for strategy in strategies:
        try:
            tables = page.extract_tables(strategy)
            for table in tables:
                # Clean the table data
                cleaned_table = []
                for row in table:
                    cleaned_row = [str(cell).strip() if cell is not None else "" for cell in row]
                    # Skip empty rows
                    if any(cell for cell in cleaned_row):
                        cleaned_table.append(cleaned_row)
                
                # Skip empty tables
                if not cleaned_table:
                    continue
                    
                # Create a key to detect duplicates
                table_key = tuple(tuple(row) for row in cleaned_table)
                if table_key not in seen_tables:
                    seen_tables.add(table_key)
                    # Format as markdown table
                    markdown_table = format_table_as_markdown(cleaned_table)
                    if markdown_table:
                        all_tables.append("```table\n" + markdown_table + "\n```")
                        
        except Exception as e:
            logger.warning(f"Error extracting table with strategy {strategy}: {str(e)}")
            continue
    
    # Extract text that might contain lists or tabular data
    text = page.extract_text() or ""
    
    # Look for common insurance sections that might contain coverage details
    coverage_terms = [
        r'(?i)(?:what(?:\'s| is) (?:not )?covered|coverage details?|benefits?|exclusions?|limitations?|\bnot covered\b|\bexcluded\b|\bnot included\b)',
        r'(?i)(?:table of benefits|schedule of benefits|summary of benefits|coverage summary)',
        r'(?i)(?:in[- ]?patient|out[- ]?patient|surgical|hospital|medical|procedure|treatment|surgery|therapy)'
    ]
    
    has_coverage_terms = any(re.search(term, text) for term in coverage_terms)
    
    # If we found coverage-related terms, be more aggressive in extracting lists
    if has_coverage_terms or not all_tables:
        # Look for bullet points and numbered lists
        list_patterns = [
            (r'(?m)^\s*[•◦‣⁃-]\s+(.+)$', '- '),  # Bullet points
            (r'(?m)^\s*\d+\.\s+(.+)$', '1. '),  # Numbered lists
            (r'(?m)^\s*[a-z]\)\s+(.+)$', 'a. '),  # Lettered lists
            (r'(?m)^\s*[ivx]+\.\s+(.+)$', 'i. '),  # Roman numerals
            (r'(?i)(?:covered|included|excluded|not covered|benefit|limitation)[:.]?\s*\n(?:\s*[-•*]\s*.+\n?)+', '- ')  # List after coverage terms
        ]
        
        for pattern, prefix in list_patterns:
            matches = re.finditer(pattern, text, re.MULTILINE | re.IGNORECASE)
            for match in matches:
                item = match.group(1).strip() if len(match.groups()) > 0 else match.group(0).strip()
                if item and len(item) > 5:  # Minimum length to avoid noise
                    # Clean up the item
                    item = re.sub(r'^\s*[\d•\-*]+[.)]?\s*', '', item)
                    item = ' '.join(item.split())
                    all_tables.append(f"```list\n{prefix}{item}\n```")
    
    # If we still don't have any tables, try to extract tabular data from text
    if not all_tables and len(text) > 100:  # Only if we have substantial text
        # Look for text that looks like a table (multiple lines with consistent spacing)
        lines = text.split('\n')
        potential_table = []
        
        for line in lines:
            # Check if line has multiple words with consistent spacing
            if len(re.findall(r'\s{2,}', line)) >= 2:
                potential_table.append(line)
            elif potential_table:
                # If we were building a table but this line doesn't match, process what we have
                if len(potential_table) >= 2:  # Need at least header + one row
                    table_text = '\n'.join(potential_table)
                    all_tables.append(f"```text_table\n{table_text}\n```")
                potential_table = []
    
    return '\n\n'.join(all_tables) if all_tables else ""

def extract_policy_details(text: str) -> Dict[str, str]:
    """
    Extract key policy details from text using regex patterns.
    
    Args:
        text: The text content to search for policy details
        
    Returns:
        Dict containing extracted policy details
    """
    details = {}
    
    # Common patterns for policy details
    patterns = {
        'sum_insured': r'(?:sum(?:\s+insured)?|cover(?:age)?(?:\s+amount)?)[:\s]+([$₹€£]?\s*\d{1,3}(?:,\d{3})*(?:\.\d{2})?)',
        'policy_number': r'(?:policy|certificate)\s*(?:no\.?|number)[:\s]+([A-Z0-9-]+)',
        'policy_period': r'(?:period of insurance|policy period)[:\s]+([^\n]+?)(?=\n\s*\n|$)',
        'premium_amount': r'(?:premium|premium amount)[:\s]+([$₹€£]?\s*\d{1,3}(?:,\d{3})*(?:\.\d{2})?)',
        'insured_name': r'(?:insured|policyholder|name of insured)[:\s]+([^\n]+?)(?=\n|$)'
    }
    
    for key, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            details[key] = match.group(1).strip()
    
    return details

def extract_text_from_pdf(file_path: str, max_pages: int = 20, timeout: int = 60) -> str:
    """
    Extract text from a PDF file with improved table and text handling.
    
    This function extracts both regular text and tabular data from PDFs,
    combining them in a way that preserves the document's structure and
    makes the content more useful for retrieval and analysis.
    
    Args:
        file_path: Path to the PDF file
        max_pages: Maximum number of pages to process (default: 20 to prevent timeouts)
        timeout: Maximum time in seconds to spend processing the PDF (default: 60s)
        
    Returns:
        str: Extracted text from the PDF, including tabular data in markdown format
        
    Raises:
        ImportError: If required dependencies are missing
        TimeoutError: If processing takes longer than the specified timeout
        ValueError: If the PDF cannot be processed
    """
    if not pdfplumber:
        raise ImportError("pdfplumber is required for PDF text extraction")
        
    import signal
    from functools import partial
    
    class ProcessingTimeout(Exception):
        pass
        
    def handler(signum, frame, file_path):
        logger.warning(f"Processing of {file_path} timed out after {timeout} seconds")
        raise ProcessingTimeout(f"PDF processing timed out after {timeout} seconds")
    
    import time
    import signal
    
    # Define a timeout handler
    class ProcessingTimeoutError(Exception):
        pass
        
    def handler(signum, frame):
        raise ProcessingTimeoutError(f"PDF processing timed out after {timeout} seconds")
    
    # Store the original signal handler
    original_handler = signal.getsignal(signal.SIGALRM)
    
    # Set the signal handler with file_path in the handler
    signal.signal(signal.SIGALRM, lambda s, f: handler(s, f, file_path))
    signal.alarm(timeout)
    
    try:
        start_time = time.time()
        full_text = []
        processed_pages = 0
        
        logger.info(f"Starting PDF text extraction from: {file_path}")
        logger.info(f"Timeout set to {timeout} seconds, max pages: {max_pages}")
        
        try:
            with pdfplumber.open(file_path) as pdf:
                total_pages = min(len(pdf.pages), max_pages) if max_pages > 0 else len(pdf.pages)
                logger.info(f"PDF contains {total_pages} pages")
                
                for i, page in enumerate(pdf.pages[:total_pages]):
                    if max_pages > 0 and processed_pages >= max_pages:
                        logger.info(f"Reached maximum page limit of {max_pages} pages")
                        break
                        
                    try:
                        page_num = i + 1
                        logger.debug(f"Processing page {page_num}/{total_pages}")
                        
                        # Check if we're approaching the timeout
                        if time.time() - start_time > timeout * 0.8:  # Use 80% of timeout for processing
                            logger.warning(f"Approaching timeout, stopping after page {i}")
                            break
                            
                        # Extract regular text with error handling and timeout
                        page_text = ""
                        try:
                            with ThreadPoolExecutor(max_workers=1) as executor:
                                future = executor.submit(page.extract_text)
                                try:
                                    page_text = future.result(timeout=5) or ""  # 5s timeout per page
                                    if page_text:
                                        page_text = clean_text(page_text)
                                except TimeoutError:
                                    logger.warning(f"Text extraction timed out on page {page_num}")
                                    future.cancel()
                        except Exception as e:
                            logger.warning(f"Error extracting text from page {page_num}: {str(e)}")
                        
                        # Extract tables with a separate timeout
                        tables_markdown = ""
                        try:
                            with ThreadPoolExecutor(max_workers=1) as executor:
                                future = executor.submit(extract_tables_with_pdfplumber, page)
                                try:
                                    tables_markdown = future.result(timeout=10) or ""  # 10s timeout for tables
                                except TimeoutError:
                                    logger.warning(f"Table extraction timed out on page {page_num}")
                                    future.cancel()
                        except Exception as e:
                            logger.warning(f"Error extracting tables from page {page_num}: {str(e)}")
                        
                        # Only add page if we have content
                        if page_text or tables_markdown:
                            page_content = []
                            page_content.append(f"\n--- Page {page_num} ---\n")
                            
                            if page_text:
                                page_content.append(page_text)
                            if tables_markdown:
                                page_content.append(tables_markdown)
                            
                            full_text.append("\n".join(page_content))
                            processed_pages += 1
                        
                        # Log progress more frequently for better monitoring
                        if (i + 1) % 5 == 0 or (i + 1) == total_pages or (i + 1) == 1:
                            elapsed = time.time() - start_time
                            remaining = max(0, timeout - elapsed)
                            logger.info(
                                f"Processed {i+1}/{total_pages} pages "
                                f"({(i+1)/total_pages:.0%}) - "
                                f"Elapsed: {elapsed:.1f}s, "
                                f"Remaining: {remaining:.1f}s"
                            )
                            
                            # Check if we should continue based on remaining time
                            if remaining < 5 and i + 1 < total_pages:
                                logger.warning(f"Less than 5 seconds remaining, stopping after page {i+1}")
                                break
                            
                    except Exception as e:
                        logger.error(f"Error processing page {page_num}: {str(e)}", exc_info=True)
                        continue
                        
        except Exception as e:
            logger.error(f"Error opening or processing PDF {file_path}: {str(e)}", exc_info=True)
            raise
        
        # Combine all pages with proper spacing
        result = "\n\n".join(full_text) if full_text else ""
        
        # Log completion
        elapsed = time.time() - start_time
        logger.info(f"Completed PDF extraction in {elapsed:.1f} seconds. Processed {processed_pages} pages.")
        
        if not result.strip():
            logger.warning("No content was extracted from the PDF")
            
        return result
        
    except ProcessingTimeoutError as te:
        logger.error(str(te))
        raise TimeoutError(str(te))
    except Exception as e:
        logger.error(f"Critical error extracting text from PDF {file_path}: {str(e)}", exc_info=True)
        raise
    finally:
        # Restore the original signal handler and disable the alarm
        signal.signal(signal.SIGALRM, original_handler)
        signal.alarm(0)

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        str: Extracted text from the DOCX
    """
    if not DocxDocument:
        raise ImportError("python-docx is required to process DOCX files.")
    
    try:
        doc = DocxDocument(file_path)
        return "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
        return ""

def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from a file based on its extension.
    
    Args:
        file_path: Path to the file
        
    Returns:
        str: Extracted text from the file
        
    Raises:
        ValueError: If the file format is not supported
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    elif file_ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file format: {file_ext}")

def clean_text(text: str) -> str:
    """
    Clean and normalize the extracted text.
    
    Args:
        text: The text to clean
        
    Returns:
        str: The cleaned text
    """
    if not text:
        return ""
    
    # Replace multiple whitespace characters with a single space
    import re
    text = re.sub(r'\s+', ' ', text)
    
    # Remove control characters
    text = ''.join(char for char in text if char.isprintable() or char.isspace())
    
    return text.strip()

def chunk_text(
    text: str, 
    chunk_size: int = 1000, 
    chunk_overlap: int = 200,
    **kwargs
) -> List[LangchainDocument]:
    """
    Split text into chunks with overlap.
    
    Args:
        text: The text to split
        chunk_size: Maximum size of each chunk
        chunk_overlap: Number of characters to overlap between chunks
        **kwargs: Additional arguments to pass to the text splitter
        
    Returns:
        List[LangchainDocument]: List of document chunks
    """
    if not text:
        return []
    
    # Initialize the text splitter
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        **kwargs
    )
    
    # Split the text into chunks
    chunks = text_splitter.split_text(text)
    
    # Convert chunks to LangchainDocument objects
    return [LangchainDocument(page_content=chunk) for chunk in chunks if chunk.strip()]

def _load_excel_with_tables(self, file_path: str) -> List[LangchainDocument]:
    """Load Excel files and convert each sheet to markdown tables."""
    try:
        import pandas as pd
        
        # Read all sheets
        xls = pd.ExcelFile(file_path)
        documents = []
        
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            
            # Clean the DataFrame
            df = df.dropna(how='all').dropna(axis=1, how='all')
            
            if not df.empty:
                # Convert to markdown
                md_content = f"# {sheet_name}\n\n{df.to_markdown(index=False)}"
                
                documents.append(LangchainDocument(
                    page_content=md_content,
                    metadata={
                        "source": f"{file_path} (Sheet: {sheet_name})",
                        "content_type": "table"
                    }
                ))
        
        return documents
        
    except Exception as e:
        logger.error(f"Error loading Excel file {file_path}: {str(e)}")
        # Fallback to unstructured loader
        loader = UnstructuredExcelLoader(file_path)
        return loader.load()

def load_document(self, file_path: str) -> List[LangchainDocument]:
    """
    Load a document and return its text content with enhanced table extraction.
    For PDFs, extracts both text and tables, and combines them intelligently.
    """
    file_path = str(Path(file_path).absolute())
    file_ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_ext == '.pdf':
            # Enhanced PDF loading with table extraction
            return self._load_enhanced_pdf(file_path)
        elif file_ext == '.txt':
            loader = TextLoader(file_path)
            return loader.load()
        elif file_ext in ['.doc', '.docx']:
            loader = UnstructuredWordDocumentLoader(file_path)
            return loader.load()
        elif file_ext == '.csv':
            loader = CSVLoader(file_path)
            return loader.load()
        elif file_ext in ['.xls', '.xlsx']:
            # For Excel files, convert to markdown tables
            return self._load_excel_with_tables(file_path)
        else:
            # Try unstructured loader as fallback
            loader = UnstructuredFileLoader(file_path)
            return loader.load()
                
    except Exception as e:
        logger.error(f"Error loading document {file_path}: {str(e)}")
        # Fallback to simple text extraction
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            return [LangchainDocument(page_content=text, metadata={"source": file_path})]
        except Exception as e2:
            logger.error(f"Fallback loading also failed for {file_path}: {str(e2)}")
            raise RuntimeError(f"Failed to load document {file_path}")

class MarkdownAwareTextSplitter(RecursiveCharacterTextSplitter):
    def split_text(self, text: str) -> List[str]:
        # Split by markdown table boundaries first
        parts = []
        current = []
        in_table = False
        table_content = []
        
        for line in text.split('\n'):
            if line.strip() == '```table':
                if in_table:
                    # End of table, add to parts
                    table_content.append(line)
                    parts.append('\n'.join(table_content))
                    table_content = []
                    in_table = False
                else:
                    # Start of table, save current content
                    if current:
                        parts.append('\n'.join(current))
                        current = []
                    table_content.append(line)
                    in_table = True
            elif in_table:
                table_content.append(line)
            else:
                current.append(line)
        
        # Add any remaining content
        if current:
            parts.append('\n'.join(current))
        
        # Now split each non-table part using the parent class
        final_chunks = []
        for part in parts:
            if part.strip().startswith('```table'):
                # Keep tables as single chunks
                final_chunks.append(part)
            else:
                # Split regular text
                final_chunks.extend(super().split_text(part))
        
        return final_chunks


def process_document(
    file_path: Union[str, os.PathLike],
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    metadata: Optional[Dict[str, Any]] = None,
    **kwargs
) -> List[LangchainDocument]:
    """
    Process a document file into chunks of text with metadata.
    
    Args:
        file_path: Path to the document file
        chunk_size: Maximum size of each chunk in characters
        chunk_overlap: Number of characters to overlap between chunks
        metadata: Additional metadata to include with each document chunk
        **kwargs: Additional arguments to pass to the text splitter
        
    Returns:
        List[LangchainDocument]: List of document chunks with metadata
    """
    try:
        # Convert PathLike to string if needed
        file_path = str(file_path) if isinstance(file_path, str) else file_path
        
        # Verify file exists
        if not Path(file_path).exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        logger.info(f"Processing document: {Path(file_path).name}")
        
        # Create default metadata if none provided
        if metadata is None:
            metadata = {}
            
        # Add file information to metadata
        if 'source' not in metadata:
            metadata['source'] = file_path
            
        # Process the document based on file type
        docs = []
        if file_path.lower().endswith('.pdf'):
            # Use the enhanced PDF loader
            loader = DocumentLoader()
            
            # Extract text with page numbers
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        # Get page text
                        page_text = page.extract_text() or ''
                        
                        # Create metadata with page number
                        page_metadata = metadata.copy()
                        page_metadata['page'] = page_num
                        
                        # Add to documents
                        docs.append(LangchainDocument(
                            page_content=page_text,
                            metadata=page_metadata
                        ))
                    except Exception as e:
                        logger.warning(f"Error processing page {page_num}: {str(e)}")
                        continue
            
            # If no pages were processed, fall back to the original method
            if not docs:
                logger.warning("No pages processed with pdfplumber, falling back to default loader")
                docs = loader.load_document(file_path)
        else:
            # For other file types, use the appropriate loader
            loader = DocumentLoader()
            docs = loader.load_document(str(file_path))
        
        # Combine all extracted text for policy details extraction
        page_texts = [doc.page_content for doc in docs]
        full_text = '\n\n'.join(page_texts)
    
        # Extract and prepend policy details
        policy_details = extract_policy_details(full_text)
        if policy_details:
            details_text = '## Policy Details\n\n' + '\n'.join(
                f'- **{key.replace("_", " ").title()}**: {value}'
                for key, value in policy_details.items()
            )
            full_text = f"{details_text}\n\n{full_text}"
    
        # Clean the text
        text = clean_text(full_text)
        
        # Create default metadata if none provided
        if metadata is None:
            metadata = {}
            
        # Ensure file_path is a Path object and get file extension
        file_path_obj = Path(file_path)
        file_ext = file_path_obj.suffix.lower()
        
        # Add file information to metadata
        metadata.update({
            'source': str(file_path_obj.name),
            'file_path': str(file_path_obj.absolute()),
            'file_name': file_path_obj.name,
            'file_extension': file_ext.lstrip('.'),
            'file_size': file_path_obj.stat().st_size,
            'is_insurance_document': is_insurance_document(text)
        })
        
        # Initialize the text splitter
        text_splitter = MarkdownAwareTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            **kwargs
        )
        
        # Split the text into chunks
        chunks = text_splitter.create_documents([text])
        
        # Add metadata to each chunk
        for i, chunk in enumerate(chunks):
            chunk.metadata.update(metadata)
            chunk.metadata['chunk_id'] = i
            chunk.metadata['total_chunks'] = len(chunks)
            
        return chunks
        
    except Exception as e:
        logger.error(f"Error processing document {file_path}: {str(e)}", exc_info=True)
        raise
